from __future__ import annotations

import json
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from eu_comply_api.config import Settings


@dataclass(slots=True)
class ParsedDocument:
    parser_name: str
    text: str


@dataclass(slots=True)
class ChunkPayload:
    chunk_index: int
    text_content: str
    char_start: int
    char_end: int
    metadata_json: dict[str, Any]


@dataclass(slots=True)
class ExtractedFactCandidate:
    field_path: str
    value: str | bool | int | float | list[str] | dict[str, str] | None
    confidence: float
    extraction_method: str
    rationale: str
    source_chunk_indexes: list[int]


class DocumentIntelligenceService:
    def __init__(self, settings: Settings) -> None:
        self._settings = settings

    def parse_document(self, filename: str, content_type: str, payload: bytes) -> ParsedDocument:
        extension = Path(filename).suffix.lower()

        if extension in {".txt", ".md"} or content_type.startswith("text/"):
            return ParsedDocument(
                parser_name="plain_text",
                text=payload.decode("utf-8", errors="ignore"),
            )
        if extension == ".json" or content_type == "application/json":
            raw = json.loads(payload.decode("utf-8"))
            return ParsedDocument(
                parser_name="json",
                text=json.dumps(raw, indent=2, sort_keys=True),
            )
        if extension == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(payload))
            pages = [page.extract_text() or "" for page in reader.pages]
            return ParsedDocument(parser_name="pdf", text="\n\n".join(pages))
        if extension == ".docx":
            from docx import Document

            document = Document(BytesIO(payload))
            paragraphs = [
                paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
            ]
            return ParsedDocument(parser_name="docx", text="\n\n".join(paragraphs))
        if extension == ".xlsx":
            from openpyxl import load_workbook

            workbook = load_workbook(filename=BytesIO(payload), read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in workbook.worksheets:
                lines.append(f"[Sheet] {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(cell) for cell in row if cell is not None]
                    if values:
                        lines.append(" | ".join(values))
            return ParsedDocument(parser_name="xlsx", text="\n".join(lines))

        raise ValueError(
            f"Unsupported artifact type for '{filename}' "
            f"({content_type or 'unknown content type'})."
        )

    def chunk_text(self, text: str) -> list[ChunkPayload]:
        normalized = text.strip()
        if not normalized:
            return []

        chunk_size = self._settings.artifact_chunk_size
        overlap = min(self._settings.artifact_chunk_overlap, chunk_size // 2)
        start = 0
        index = 0
        chunks: list[ChunkPayload] = []
        while start < len(normalized):
            end = min(start + chunk_size, len(normalized))
            chunk_text = normalized[start:end].strip()
            if chunk_text:
                chunks.append(
                    ChunkPayload(
                        chunk_index=index,
                        text_content=chunk_text,
                        char_start=start,
                        char_end=end,
                        metadata_json={"length": len(chunk_text)},
                    )
                )
                index += 1
            if end >= len(normalized):
                break
            start = max(end - overlap, start + 1)
        return chunks

    def extract_fact_candidates(self, chunks: list[ChunkPayload]) -> list[ExtractedFactCandidate]:
        results: list[ExtractedFactCandidate] = []
        for chunk in chunks:
            lower_text = chunk.text_content.lower()

            if any(
                keyword in lower_text
                for keyword in ["hiring", "recruit", "candidate", "employment"]
            ):
                results.append(
                    ExtractedFactCandidate(
                        field_path="use_case.domain",
                        value="employment",
                        confidence=0.72,
                        extraction_method="heuristic_keyword",
                        rationale="Employment-related keywords were detected in the artifact text.",
                        source_chunk_indexes=[chunk.chunk_index],
                    )
                )

            activity_values: list[str] = []
            if "screening" in lower_text or "candidate screening" in lower_text:
                activity_values.append("screening")
            if "recruit" in lower_text or "hiring" in lower_text:
                activity_values.append("recruitment")
            if "promotion" in lower_text:
                activity_values.append("promotion")
            if activity_values:
                results.append(
                    ExtractedFactCandidate(
                        field_path="use_case.activities",
                        value=sorted(set(activity_values)),
                        confidence=0.77,
                        extraction_method="heuristic_keyword",
                        rationale=(
                            "The document includes activity keywords that map to Annex III-style "
                            "employment decision support activities."
                        ),
                        source_chunk_indexes=[chunk.chunk_index],
                    )
                )

            if any(keyword in lower_text for keyword in ["credit", "loan", "insurance", "bank"]):
                results.append(
                    ExtractedFactCandidate(
                        field_path="use_case.domain",
                        value="finance",
                        confidence=0.72,
                        extraction_method="heuristic_keyword",
                        rationale="Finance-related keywords were detected in the artifact text.",
                        source_chunk_indexes=[chunk.chunk_index],
                    )
                )

            if "chatbot" in lower_text or "conversational ai" in lower_text:
                results.append(
                    ExtractedFactCandidate(
                        field_path="system.modalities",
                        value=["chatbot"],
                        confidence=0.79,
                        extraction_method="heuristic_keyword",
                        rationale=(
                            "Conversational system keywords indicate a chatbot-style interface."
                        ),
                        source_chunk_indexes=[chunk.chunk_index],
                    )
                )

            if "remote biometric identification" in lower_text:
                results.append(
                    ExtractedFactCandidate(
                        field_path="system.capabilities.remote_biometric_identification",
                        value=True,
                        confidence=0.88,
                        extraction_method="heuristic_keyword",
                        rationale=(
                            "The document explicitly references remote biometric identification."
                        ),
                        source_chunk_indexes=[chunk.chunk_index],
                    )
                )

            if all(
                keyword in lower_text
                for keyword in ["law enforcement", "public space", "real-time"]
            ):
                results.append(
                    ExtractedFactCandidate(
                        field_path="deployment.context.law_enforcement_real_time_public_space",
                        value=True,
                        confidence=0.84,
                        extraction_method="heuristic_keyword",
                        rationale=(
                            "The document describes law-enforcement real-time public-space usage."
                        ),
                        source_chunk_indexes=[chunk.chunk_index],
                    )
                )

            if any(
                keyword in lower_text
                for keyword in ["large language model", "llm", "generative ai"]
            ):
                results.append(
                    ExtractedFactCandidate(
                        field_path="system.uses_generative_ai",
                        value=True,
                        confidence=0.7,
                        extraction_method="heuristic_keyword",
                        rationale=(
                            "The artifact mentions a large language model or generative AI usage."
                        ),
                        source_chunk_indexes=[chunk.chunk_index],
                    )
                )

        return self._deduplicate_candidates(results)

    def _deduplicate_candidates(
        self, candidates: list[ExtractedFactCandidate]
    ) -> list[ExtractedFactCandidate]:
        deduplicated: dict[tuple[str, str], ExtractedFactCandidate] = {}
        for candidate in candidates:
            key = (candidate.field_path, json.dumps(candidate.value, sort_keys=True))
            existing = deduplicated.get(key)
            if existing is None or existing.confidence < candidate.confidence:
                deduplicated[key] = candidate
        return list(deduplicated.values())
